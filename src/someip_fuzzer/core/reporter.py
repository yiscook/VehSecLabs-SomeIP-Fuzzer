"""报告生成引擎 — Jinja2 + WeasyPrint(PDF) + python-docx(DOCX)。

用法::

    from someip_fuzzer.core.reporter import Reporter, ReportConfig
    from someip_fuzzer.data.crash_store import CrashStorage

    store = CrashStorage("crashes.db")
    crashes = store.list_all()

    config = ReportConfig(company="VehSecLabs", author="李奇")
    session = {"id": "SESS-001", "target": "192.168.1.1:30509",
               "date": "2026-05-11", "duration": 3600, "transport": "UDP"}

    reporter = Reporter()
    reporter.to_html(session, crashes, config, Path("report.html"))
    reporter.to_pdf(session, crashes, config, Path("report.pdf"))
    reporter.to_docx(session, crashes, config, Path("report.docx"))
"""

from __future__ import annotations

import binascii
from dataclasses import dataclass, field
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from someip_fuzzer.data.crash_store import CrashRecord

_TEMPLATES_DIR = Path(__file__).parents[3] / "templates"


# ── 报告配置 ──────────────────────────────────────────────────────────────────

@dataclass
class ReportConfig:
    title: str = "SOME/IP 模糊测试安全报告"
    company: str = "VehSecLabs"
    author: str = ""
    logo_path: Path | None = None
    include_methodology: bool = True
    include_vulnerabilities: bool = True
    include_reproduction: bool = True
    include_recommendations: bool = True
    include_cvss: bool = True
    include_raw_appendix: bool = False


# ── Jinja2 过滤器 ─────────────────────────────────────────────────────────────

def _hexdump(data: bytes, bytes_per_row: int = 16) -> str:
    """将字节数据渲染为 hexdump 格式。"""
    lines = []
    for i in range(0, len(data), bytes_per_row):
        chunk = data[i:i + bytes_per_row]
        hex_part = " ".join(f"{b:02X}" for b in chunk).ljust(bytes_per_row * 3 - 1)
        asc_part = "".join(chr(b) if 32 <= b < 127 else "." for b in chunk)
        lines.append(f"{i:04X}: {hex_part}  {asc_part}")
    return "\n".join(lines)


def _format_crash_id(index: int) -> str:
    return f"CRASH-{index:03d}"


# ── Reporter 主类 ─────────────────────────────────────────────────────────────

class Reporter:
    """统一报告生成入口，支持 HTML / PDF / DOCX 三种格式。"""

    def __init__(self, templates_dir: Path | None = None) -> None:
        self._templates_dir = templates_dir or _TEMPLATES_DIR
        self._env = self._build_jinja_env()

    def _build_jinja_env(self):
        from jinja2 import Environment, FileSystemLoader, select_autoescape
        env = Environment(
            loader=FileSystemLoader(str(self._templates_dir)),
            autoescape=select_autoescape(["html"]),
        )
        env.filters["hexdump"] = _hexdump
        env.filters["format_crash_id"] = _format_crash_id
        env.filters["split"] = lambda s, sep: s.split(sep)
        return env

    # ── 数据准备 ─────────────────────────────────────────────────────────────

    def _build_summary(self, crashes: list[CrashRecord]) -> dict:
        return {
            "crashes": len(crashes),
            "critical": sum(1 for c in crashes if c.severity == "critical"),
            "high":     sum(1 for c in crashes if c.severity == "high"),
            "medium":   sum(1 for c in crashes if c.severity == "medium"),
            "low":      sum(1 for c in crashes if c.severity == "low"),
            "sent":     0,       # 调用方可覆盖
            "duration": 0,
            "strategies": 93,   # Phase 2 实现数量
        }

    def _build_context(
        self,
        session: dict,
        crashes: list[CrashRecord],
        config: ReportConfig,
        css: str = "",
    ) -> dict:
        summary = self._build_summary(crashes)
        summary.update({k: session.get(k, summary[k]) for k in ("sent", "duration") if k in session})
        return {
            "session": session,
            "crashes": crashes,
            "summary": summary,
            "config": config,
            "css": css,
        }

    # ── HTML ─────────────────────────────────────────────────────────────────

    def render_html(
        self,
        session: dict,
        crashes: list[CrashRecord],
        config: ReportConfig,
    ) -> str:
        """渲染为 HTML 字符串。"""
        css_path = self._templates_dir / "styles" / "report.css"
        css = css_path.read_text(encoding="utf-8") if css_path.exists() else ""
        tpl = self._env.get_template("report.html")
        return tpl.render(**self._build_context(session, crashes, config, css=css))

    def to_html(
        self,
        session: dict,
        crashes: list[CrashRecord],
        config: ReportConfig,
        output: Path,
    ) -> Path:
        """生成 HTML 报告文件。"""
        output = Path(output)
        output.parent.mkdir(parents=True, exist_ok=True)
        output.write_text(self.render_html(session, crashes, config), encoding="utf-8")
        return output

    # ── PDF ──────────────────────────────────────────────────────────────────

    def to_pdf(
        self,
        session: dict,
        crashes: list[CrashRecord],
        config: ReportConfig,
        output: Path,
    ) -> Path:
        """生成 PDF 报告（WeasyPrint）。"""
        from weasyprint import HTML  # type: ignore[import]
        output = Path(output)
        output.parent.mkdir(parents=True, exist_ok=True)
        html_str = self.render_html(session, crashes, config)
        HTML(string=html_str, base_url=str(self._templates_dir)).write_pdf(str(output))
        return output

    # ── DOCX ─────────────────────────────────────────────────────────────────

    def to_docx(
        self,
        session: dict,
        crashes: list[CrashRecord],
        config: ReportConfig,
        output: Path,
    ) -> Path:
        """生成 DOCX 报告（python-docx）。"""
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        output = Path(output)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()

        # 设置默认字体
        doc.styles["Normal"].font.name = "Microsoft YaHei"
        doc.styles["Normal"].font.size = Pt(11)

        # 封面
        cover = doc.add_heading(config.title, level=0)
        cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"{config.company}\n{config.author}\n{session.get('date', '')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        # 执行摘要
        summary = self._build_summary(crashes)
        doc.add_heading("1. 执行摘要", level=1)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Light Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "指标"
        hdr[1].text = "数值"
        for label, val in [
            ("靶机地址", session.get("target", "—")),
            ("发包总量", str(summary["sent"])),
            ("崩溃总数", str(summary["crashes"])),
            ("严重", str(summary["critical"])),
            ("高危", str(summary["high"])),
            ("中危", str(summary["medium"])),
        ]:
            row = tbl.add_row().cells
            row[0].text = label
            row[1].text = val

        # 漏洞详情
        if config.include_vulnerabilities and crashes:
            doc.add_page_break()
            doc.add_heading("2. 漏洞详情", level=1)
            for i, crash in enumerate(crashes, 1):
                doc.add_heading(f"{i}. {crash.mutator_name}  (CVSS {crash.cvss_score:.1f})", level=2)
                p = doc.add_paragraph()
                p.add_run("严重度：").bold = True
                p.add_run(crash.severity.upper())
                p.add_run("  | 检测方式：").bold = True
                p.add_run(crash.detection_method)
                p.add_run("  | 时间：").bold = True
                p.add_run(crash.timestamp[:19])

                if config.include_reproduction:
                    doc.add_paragraph("触发报文（Hex）：").runs[0].bold = True
                    hex_para = doc.add_paragraph(_hexdump(crash.triggering_bytes))
                    hex_para.style = "No Spacing"
                    for run in hex_para.runs:
                        run.font.name = "Consolas"
                        run.font.size = Pt(9)

                if config.include_recommendations:
                    rec = doc.add_paragraph()
                    rec.add_run("修复建议：").bold = True
                    rec.add_run(f"对 {crash.mutator_name.split('.')[0]} 相关字段添加严格的长度/范围校验。")

        # 页脚
        section = doc.sections[0]
        footer = section.footer
        footer.paragraphs[0].text = f"本报告由 VehSecLabs-SomeIP-Fuzzer 自动生成 | {config.company}"

        doc.save(str(output))
        return output
